from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "MateClaw智能排障-作品方案.docx"
ASSETS = ROOT / "assets"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
ACCENT = "DD704B"
GREEN = "2F7265"
MUTED = "6F665F"
PALE = "F4F6F9"
WARM = "F8EEE7"
GREEN_PALE = "EAF4EF"
WHITE = "FFFFFF"
BLACK = "221E1B"
DOC_FONT = "Arial Unicode MS"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_header(section, left: str, right: str) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)
    table.cell(0, 0).text = left
    table.cell(0, 1).text = right
    for idx, cell in enumerate(table.rows[0].cells):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
        for run in cell.paragraphs[0].runs:
            run.font.name = DOC_FONT
            run.font.size = Pt(8)
            run.font.color.rgb = rgb(MUTED)
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run2 = paragraph.add_run(" 页")
    for current in (run, run2):
        current.font.name = DOC_FONT
        current.font.size = Pt(8)
        current.font.color.rgb = rgb(MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_repeat_header(section, "MateClaw 智能排障", "作品方案 · 2026-08-11")
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = DOC_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = DOC_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = DOC_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    props = doc.core_properties
    props.title = "MateClaw 智能排障作品方案"
    props.subject = "场景描述、AI 工具使用、效能提升、Demo 与投产计划"
    props.author = "MateClaw 项目组"
    props.keywords = "MateClaw, 智能排障, Guance, 只读取证, 确定性判断"


def add_text(doc, text: str, *, bold=False, color=None, size=None, align=None, after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = rgb(color)
    if size:
        run.font.size = Pt(size)
    run.font.name = DOC_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    run._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    return p


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_callout(doc, title: str, body: str, fill=WARM, accent=ACCENT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = rgb(accent)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths=None, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for idx, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(width)
        set_cell_shading(cell, PALE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(font_size)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for idx, (value, width) in enumerate(zip(row_values, widths)):
            cell = cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.name = DOC_FONT
            r.font.size = Pt(font_size)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
            r._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, image_name: str, caption: str, width=6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(ASSETS / image_name), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(10)
    r = cp.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(MUTED)


def add_page_break(doc) -> None:
    doc.add_page_break()


def build() -> None:
    doc = Document()
    configure_document(doc)

    # Cover: proposal_centerpiece pattern with centered title stack + two-column metadata.
    add_text(doc, "MateClaw 项目组", bold=True, color=MUTED, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_text(doc, "MateClaw 智能排障", bold=True, color=BLACK, size=26,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    add_text(doc, "让排障能力不再只掌握在少数熟手手里", color=ACCENT, size=15,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_text(doc, "作品方案｜场景描述 · AI 工具使用 · 效能提升 · Demo 与投产计划",
             bold=True, color=MUTED, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    meta.autofit = False
    left = [
        ("作品类型", "IT 智能排障平台"),
        ("首个场景", "CSDP / csdp-wechat / 904003"),
        ("证据来源", "真实 Guance 只读数据"),
        ("当前阶段", "真实单场景竖线已贯通"),
    ]
    right = [
        ("核心方法", "三次取证 + 成功对照 + 确定性判断"),
        ("真实排障单", "diag-acee292e…ec2e"),
        ("安全边界", "零生产写；错误码路径零模型调用"),
        ("版本日期", "2026-08-11"),
    ]
    for i in range(4):
        for j, pair in enumerate((left[i], right[i])):
            cell = meta.cell(i, j)
            set_cell_margins(cell, top=110, start=150, bottom=110, end=150)
            if i % 2 == 0:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(pair[0])
            r.bold = True
            r.font.color.rgb = rgb(BLUE)
            p2 = cell.add_paragraph(pair[1])
            p2.paragraph_format.space_after = Pt(0)
    add_text(doc, "", after=8)
    add_callout(
        doc,
        "作品主张",
        "MateClaw 让二线能够基于平台开展标准化排障，让三线开发快速进入陌生系统的有效调查，并在重大故障中帮助开发团队基于同一份证据快速收敛问题范围。",
        fill=GREEN_PALE,
        accent=GREEN,
    )
    add_text(doc, "本材料只陈述已经由当前代码、真实排障单和验证记录支持的能力；未完成项在第七章单列。",
             color=MUTED, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    add_page_break(doc)
    doc.add_heading("作品摘要", level=1)
    add_text(doc, "MateClaw 智能排障首先解决三个真实问题：二线看到告警却不知道下一步怎么查；三线开发接手陌生系统时过度依赖少数熟手；重大故障中多人重复查询、判断口径不一，问题范围迟迟无法收敛。平台把已经验证的调查方法、真实证据和人工责任边界放进同一张排障单。")
    add_callout(doc, "已验证的真实结果", "CSDP 的 csdp-wechat 服务在 ITGW 904003 故障窗口中，2 个失败请求均出现“内容拦截”，36 个正常请求均未出现；本次系统调查阶段记录约 4 秒。", fill=WARM)

    doc.add_heading("核心能力一览", level=2)
    add_table(
        doc,
        ["能力", "当前实现", "用户价值"],
        [
            ("告警结构化", "保存系统、服务、时间、错误码和影响线索", "减少补问和范围误判"),
            ("三次只读取证", "失败检索 → 关联日志 → 正常请求对照", "自动串联调查上下文"),
            ("确定性判断", "按冻结规则版本复算条件", "结论可解释、可重放"),
            ("双视图", "业务结果 + 开发证据", "同一事实服务不同角色"),
            ("知识闭环", "结案生成候选，回放和审核后启用", "把个人经验变成组织资产"),
        ],
        widths=[1.35, 2.55, 2.6],
    )
    doc.add_heading("当前边界", level=2)
    add_bullet(doc, "已经证明一条真实 Guance 错误码场景可以端到端运行。")
    add_bullet(doc, "尚未证明全部系统已经接入，也未达到正式 20–30 条历史样本批次验收。")
    add_bullet(doc, "平台不执行生产变更；确认结论只推进排障单状态。")

    doc.add_heading("一、场景描述", level=1)
    doc.add_heading("1.1 项目要解决的三个真实问题", level=2)
    add_table(
        doc,
        ["使用者 / 场景", "现在遇到的问题", "平台首先要交付的结果"],
        [
            ("二线", "看到告警现象，但不知道该查什么、如何判断", "按已审核方法启动并推进只读调查；证据不足时说清缺什么、为什么升级"),
            ("三线开发", "接手陌生系统时不熟悉字段、关联 ID 和历史经验，排障依赖少数熟手", "先获得清晰调查路径、证据和结论，降低对系统熟悉度与个人经验的依赖"),
            ("重大故障团队", "多人重复查询、信息口径不一，异常与故障的关系难以确认", "统一证据、判断口径和责任边界，快速收敛问题范围并减少无效沟通"),
        ],
        widths=[1.25, 2.6, 2.65],
        font_size=8.8,
    )
    add_callout(doc, "不是降低开发要求", "平台降低的是排障对系统熟悉度、少数专家和个人记忆的依赖；真实字段核对、最终因果确认和生产处置仍由授权人员负责。", fill=PALE, accent=BLUE)

    doc.add_heading("1.2 三个问题与三种处理方式不能混为一谈", level=2)
    add_text(doc, "三个问题说的是作品要交付的业务结果；三种处理方式说的是平台面对不同成熟度场景时的工作方式。")
    add_table(
        doc,
        ["场景成熟度", "平台怎么做", "当前状态"],
        [
            ("已有审核方法", "直接执行冻结查询和确定性条件", "ITGW 904003 真实案例已跑通"),
            ("新问题已由人解决", "结案后形成候选方法，经回放和人工审核后再复用", "候选、回放和审核机制已实现"),
            ("完全没有方法", "只允许受限、只读的探索，无证据即停止", "通用自主规划尚未通过投产验证"),
        ],
        widths=[1.45, 3.25, 1.8],
        font_size=8.8,
    )

    doc.add_heading("1.3 排障过程中的共性断点", level=2)
    for item in (
        "入口信息不完整：告警有系统、服务和现象，但缺少稳定调查路径。",
        "证据分散：失败日志、同一请求的关联日志以及正常请求分布在不同查询中。",
        "判断不可复核：看见某条异常日志，并不能证明它只集中在失败请求中。",
        "经验难沉淀：查询条件、判断标准和停止原因没有变成可复用资产。",
    ):
        add_number(doc, item)
    doc.add_heading("1.4 真实案例", level=2)
    add_table(
        doc,
        ["字段", "内容"],
        [
            ("事件", "客服数字化（WECHAT）— ITGW访问失败【904003】"),
            ("时间", "2026-08-07 17:12（Asia/Shanghai）"),
            ("系统 / 服务", "CSDP / csdp-wechat"),
            ("集群线索", "sz3-s-k8s；仅作为资产元数据，不猜测观测云字段"),
            ("真实排障单", "diag-acee292ecd7647288e2c39e80007ec2e"),
            ("结果", "已定位、待人工确认；fixtureMode=false"),
        ],
        widths=[1.5, 5.0],
        font_size=9.8,
    )
    add_text(doc, "系统先查失败请求并取得真实关联 ID，再读取同一 ID 下的关联日志，最后在同一时间窗口选择正常请求做对照。两条确定性判断条件成立后，规则把异常环节定位到 ITGW 内容安全策略拦截。")
    add_callout(doc, "为什么仍然“待确认”", "证据支持异常特征与故障有关，但生产处置与最终因果确认仍由授权负责人完成；平台不会把相关性包装成自动修复。", fill=PALE, accent=BLUE)

    doc.add_heading("1.5 谁负责接入，开发怎么开始", level=2)
    add_text(doc, "接入分成一次性准备与日常排障两条路径。平台管理员统一配置数据源，系统负责人维护本系统观测资产与查询规则；开发和值班人员日常只需要从告警创建排障单，不必自己编写 DQL。")
    add_table(
        doc,
        ["角色", "操作入口", "主要职责"],
        [
            ("平台管理员", "取证接入 → 数据源联调", "配置并验证 Guance 连接、认证、超时和只读边界；同一数据源统一配置"),
            ("系统负责人", "系统与模块、查询规则", "登记系统、服务和环境，维护失败、关联日志与正常对照规则"),
            ("审核负责人", "排障规则库、路由与绑定", "完成只读试跑、回放、规则审核和版本启用"),
            ("开发 / 值班", "智能排障 → 排障工作台", "从告警创建排障单，查看证据，确认、驳回或转人工"),
        ],
        widths=[1.15, 2.0, 3.35],
        font_size=8.8,
    )
    doc.add_heading("开发的五步使用路径", level=3)
    for index, item in enumerate((
        "收到告警后进入排障工作台，点击创建排障单。",
        "填写系统、服务、精确告警时间、错误码和故障现象。",
        "系统自动选择已审核方法并执行只读取证，开发不需要编写 DQL。",
        "详情页先看结论，需要复核时再看关键节点、完整过程和证据关系。",
        "负责人确认或驳回，在平台外完成生产处置并登记结果。",
    ), start=1):
        paragraph = add_text(doc, f"{index}. {item}", after=4)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    add_callout(doc, "无法自动调查时怎么办", "系统未接入、规则未审核或真实证据不可用时，平台会明确显示配置缺口或“证据不足”，并转回人工排障；不会用 Demo、历史猜测或模型回答替代真实证据。", fill=PALE, accent=BLUE)

    doc.add_heading("二、方案设计", level=1)
    doc.add_heading("2.1 七阶段主流程", level=2)
    add_table(
        doc,
        ["阶段", "系统要回答的问题", "本案例"],
        [
            ("1. 收到告警", "哪个系统、哪个服务、何时、什么错误？", "CSDP / csdp-wechat / 904003"),
            ("2. 选择方法", "是否有审核通过的调查方法？", "冻结 904003 排障方法 v2"),
            ("3. 明确证据", "需要哪些事实才能判断？", "失败、关联日志、正常对照"),
            ("4. 连接数据源", "由哪个只读适配器执行？", "Guance 三份绑定"),
            ("5. 获取证据", "真实查询返回了什么？", "3 份规范化证据"),
            ("6. 计算条件", "证据是否满足审核条件？", "2 条条件成立"),
            ("7. 形成结论", "定位、弃权还是转人工？", "已定位，待人工确认"),
        ],
        widths=[1.15, 3.0, 2.35],
        font_size=8.8,
    )
    doc.add_heading("2.2 三次只读取证", level=2)
    add_table(
        doc,
        ["次序", "查询目的", "运行时输入", "安全输出"],
        [
            ("1", "找失败请求", "服务、错误码、绝对时间窗", "失败数量、受限关联 ID"),
            ("2", "查关联日志", "第一步返回的真实关联 ID", "时间、服务、级别、白名单标记"),
            ("3", "找正常请求作对照", "同一服务、同一窗口、审核特征", "两组请求数和特征出现数"),
        ],
        widths=[0.55, 1.6, 2.15, 2.2],
        font_size=9,
    )
    doc.add_heading("2.3 安全和治理", level=2)
    add_bullet(doc, "错误码已命中时全程零模型调用，结论来自冻结规则和规范化证据。")
    add_bullet(doc, "不返回或持久化 API Key、DQL、原始日志正文和业务内容。")
    add_bullet(doc, "缺少规则、数据连接或证据时按 fail-closed 输出“证据不足”。")
    add_bullet(doc, "只提供只读证据能力，不注册生产写执行器。")

    doc.add_heading("三、AI 工具使用说明", level=1)
    doc.add_heading("3.1 产品运行时", level=2)
    add_table(
        doc,
        ["使用位置", "AI 的职责", "权威边界"],
        [
            ("错误码已命中", "不调用模型", "审核规则 + 真实证据 + 确定性计算"),
            ("未知场景", "受限提出只读调查建议", "硬白名单；不可用时直接停止"),
            ("历史案例归纳", "生成排障方法草稿", "回放、确定性校验和人工审核后才启用"),
            ("证据解释", "把技术字段与计数翻译成人话", "不改写底层事实与判断结果"),
        ],
        widths=[1.3, 2.7, 2.5],
        font_size=9.2,
    )
    doc.add_heading("3.2 研发阶段", level=2)
    add_bullet(doc, "Codex：代码库理解、跨前后端实现、测试补齐、浏览器验收和交付材料生成。")
    add_bullet(doc, "Claude：早期架构讨论、方案反证和交互方案迭代。")
    add_bullet(doc, "人工：确认真实字段、判断阈值、审核启用和生产处置。")
    add_callout(doc, "验证原则", "AI 生成的代码和方案必须通过单元测试、固定回放、真实只读调用与人工复核；不能以“模型说可以”作为完成标准。", fill=GREEN_PALE, accent=GREEN)

    doc.add_heading("四、效能提升说明", level=1)
    add_table(
        doc,
        ["环节", "传统方式", "MateClaw"],
        [
            ("查询路径", "依赖个人经验，容易查错窗口或字段", "审核方法固定顺序和绝对时间窗"),
            ("关联信息", "手工复制 ID，多页面切换", "第一步结果自动成为后续唯一输入"),
            ("异常判断", "看到错误日志后凭经验推断", "失败与正常请求同窗对照后计算"),
            ("过程复核", "聊天、截图和链接分散", "关键节点、七阶段和证据关系集中"),
            ("经验沉淀", "依赖口口相传", "结案形成候选，回放审核后复用"),
        ],
        widths=[1.15, 2.55, 2.8],
        font_size=9,
    )
    add_text(doc, "当前真实记录：三次 Guance-only 查询取得规范化证据；系统调查阶段约 4 秒；2 个失败请求均出现异常特征，36 个正常请求均未出现；2 条确定性判断条件成立。")
    add_callout(doc, "4 秒证明了什么", "它只证明这条已审核方法能在真实 Guance 数据上完成一次系统调查，不等于已证明二线独立排障、三线快速上手和重大故障收敛效率。", fill=WARM)
    doc.add_heading("三个问题的效果验收", level=2)
    add_table(
        doc,
        ["要验证的问题", "真实试用方式", "主要指标"],
        [
            ("二线能基于平台排障", "让二线人员不由三线代操，独立处理已审核场景", "独立完成率、升级比例、首个可读结论耗时"),
            ("三线开发能快速上手", "让不熟悉该系统的开发完成同一批排障任务", "首次排障耗时、求助次数、证据完整度"),
            ("重大故障能更快收敛", "在真实演练或历史复盘中让多角色共用同一证据链", "范围收敛耗时、重复查询数、无效转交数"),
        ],
        widths=[1.6, 2.75, 2.15],
        font_size=8.6,
    )
    add_text(doc, "以上人员效果指标尚待真实试用。此外，仍需在 20–30 条历史样本上统计补问、系统调查和人工采纳三个阶段的 p50/p95。", color=MUTED, size=9.5)

    add_page_break(doc)
    doc.add_heading("五、Demo 与效果示意", level=1)
    doc.add_heading("5.1 结果总览", level=2)
    add_text(doc, "结果页先回答“发生了什么、当前结论是什么、谁来处置”，并把补问、系统调查和人工采纳三个耗时阶段分开记录。")
    add_figure(doc, "01-排障结果总览.png", "图 1　真实排障结果总览：已定位、待确认，平台不执行生产变更")

    doc.add_heading("5.2 四个关键节点", level=2)
    add_text(doc, "默认使用“发生了什么、按什么方法查、查到了什么、最后结论”四个大白话节点；开发需要复核时，再切换到完整七步或证据关系。")
    add_figure(doc, "02-四个关键节点.png", "图 2　四个关键节点：左侧选阶段，右侧查看本次结果和为什么重要")
    doc.add_heading("5.3 证据对照", level=2)
    add_text(doc, "页面不展示容易混淆的 2/2、0/36 缩写，而是直接说明两组请求分别发生了什么；PS ID 区域明确标注为关联日志，不冒充完整跨服务 Trace。")
    add_figure(doc, "03-证据对照.png", "图 3　失败请求与正常请求对照：异常特征只集中在失败请求中")

    doc.add_heading("六、创新点", level=1)
    for title, body in (
        ("成功样本对照", "不仅找到失败日志，还验证异常特征是否集中在失败请求中。"),
        ("证据优先于模型", "真实查询、规范化证据和确定性规则构成权威链，AI 只做建议和解释。"),
        ("同一证据双投影", "业务负责人看结论和责任边界，开发看查询、判断条件和证据引用。"),
        ("可弃权设计", "没有证据就停止，比生成一个听起来合理的答案更安全。"),
        ("经验生产闭环", "真实结案生成候选规则，但必须经回放和人工审核后才能影响下一次排障。"),
    ):
        doc.add_heading(title, level=3)
        add_text(doc, body, after=5)

    doc.add_heading("七、完成度与投产计划", level=1)
    doc.add_heading("7.1 已完成", level=2)
    add_bullet(doc, "MateClaw 内部确定性排障领域模块、正式工作台和排障单生命周期。")
    add_bullet(doc, "Guance 只读适配器、三次取证、脱敏压缩和失败/正常请求对照。")
    add_bullet(doc, "CSDP / csdp-wechat / 904003 真实案例端到端跑通。")
    add_bullet(doc, "查询目录、系统接入、排障方法审核、历史回放和案例沉淀入口。")
    doc.add_heading("7.2 投产前必须完成", level=2)
    for index, item in enumerate((
        "由各系统 owner 核对 measurement、字段、时间单位、索引和查询窗口。",
        "补齐查询规则，冻结至少 20 个真实可执行目标；当前正式目标目录仍为 0/20。",
        "对当前配置指纹完成 owner acceptance，再运行 20–30 条历史影子样本。",
        "统计三段北极星耗时和误判/弃权情况，形成 p50/p95 与准入阈值。",
        "完成 RBAC、审计、告警通道、回滚演练和生产运行手册。",
    ), start=1):
        paragraph = add_text(doc, f"{index}. {item}", after=5)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)

    add_callout(doc, "作品价值", "MateClaw 让二线能基于平台启动和推进标准化排障，让三线开发降低进入陌生系统调查的前置门槛，让重大故障团队围绕同一份证据快速收敛范围；并让一次真实排障变成下一次可审核、可重复执行的组织能力。", fill=GREEN_PALE, accent=GREEN)
    add_text(doc, "Demo 入口（需本地服务和登录态）：http://127.0.0.1:5173/troubleshooting?view=detail&diagnosisId=diag-acee292ecd7647288e2c39e80007ec2e", color=MUTED, size=8.8, after=0)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
